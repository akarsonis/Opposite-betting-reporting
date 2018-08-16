import random

import time
from datetime import datetime, timedelta
import os
import sys

from tkinter import *

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches

import xlrd
import openpyxl

from collections import Counter
import re

from bs4 import BeautifulSoup
import requests

#Text

intro = 'Ezugi Fraud and Risk department detection tools, reports and operations are based on a search for particular betting patterns to identify advantage play or cheating behavior within Ezugi gaming network. Accounts that are a subject of this report are under Ezugi Fraud and Risk department\'s observation since the players\' gameplay revealed signs of Opposite betting.'

explanatory1 = 'After detailed investigation, it was detected that the players below are playing in pairs with each other by placing the same bets on opposite betting spots. Please find suspected accounts\' information below:'
    
explanatory2_roulette = 'The players have been participating in same games, playing Roulette and covering the whole layout with outside bets. Please find examples below:'
explanatory2_baccarat = 'The players have been participating in same games, playing Baccarat wagering same amounts against each other. Please find examples below:'

conclusion = 'These players revealed signs of opposite betting. Please let us know if any further assistance is required.'

#Adding Function

lista = []

counter = 0
def list_append(event):
    global counter
    lista.append(link_entry.get())
    link_entry.delete(0, END)
    counter += 1
    link_label_var.set('BO Link  ' + str(counter))

#Main Function

def main_function(event):
    
    document = Document('Ezugi_template_Opposite players.docx')
    
    #Essential functions before main function
    
    #Adding 1st/2nd Names on the cover page
    
    if os.environ['USERNAME'] == 'Aleksandrs':
        run01 = document.paragraphs[13].add_run(' Aleksandrs Karsonis')
        font01 = run01.font
        font01.size = Pt(14)    
        if os.environ['USERNAME'] == 'Arturs':
            run02 = document.paragraphs[13].add_run(' Arturs Lusis')
            font02 = run01.font
            font02.size = Pt(14)
            if os.environ['USERNAME'] == 'Alina':
                run02 = document.paragraphs[13].add_run(' Alina Heifeca')
                font02 = run01.font
                font02.size = Pt(14)        
    
    run2 = document.paragraphs[15].add_run(time.strftime(" %d.%m.%y"))
    font2 = run2.font
    font2.size = Pt(14)
    
    #cell declaration
    
    table0 = document.tables[0] #Top Cell
    cell011 = table0.cell(1,1)
    
    table1 = document.tables[1] #Body Cell
    cell110 = table1.cell(1,0)
    
    table2 = document.tables[2] #Final cell
    cell210 = table2.cell(1,0)
    
    table1_inside_cell110 = cell110.tables[0] #Table inside cell
    
    global intro
    global explanatory1
    global explanatory2_roulette
    global explanatory2_baccarat
    global conclusion
    
    #Investigation timeframe
    
    if dates_entry.get() == '':
        run3 = document.paragraphs[14].add_run(' All history')
        font2 = run3.font
        font2.size = Pt(14)
    else:
        run3 = document.paragraphs[14].add_run(' ' + dates_entry.get())
        font2 = run3.font
        font2.size = Pt(14)    
    
    #Logging in BO
    try: 
        with requests.Session() as c:
            url = 'https://bo.livetables.io/office.php?page=login'
            USERNAME = 'akarsonis'
            PASSWORD = 'ezugi123456'
            c.get(url)
            login_data = dict(username=USERNAME, password=PASSWORD, 
                              language_view='english', submit='Login')
            c.post(url, 
                   data=login_data, 
                   headers={'Referer': 
                            'https://bo.livetables.io/office.php?page=login'})
        
        for link in lista:
            page = c.get(link)
            soup = BeautifulSoup(page.content, 'html.parser')
            
            #Parsing Operator ID
            operator_id = str(soup.find('img', {'height' : '20'}))
            operator_id = ((operator_id.split('ID: '))[1].split(' " width="20"/>')[0])
            operator_id = operator_id.replace('\n', '')
            operator_id = operator_id.replace('  ', ' ')
            operator_id = operator_id[:-1]
            
            #Parsing UID
            uid = str(soup.find('input', {'name' : 'PlayerDisplay'}))
            uid = uid.split('" type="text" value="')[1].split('"/>')[0]
            
            #Parsing screen name
            screen_name = str(soup.find('img', {'width' : '25'}))
            screen_name = screen_name.split(' - ')[1].split(' Operator: ')[0]
            screen_name = screen_name.replace('\n', '')
            
            #Parsing currency
            currency = str(soup.find('td', {'class' : 'grid_cell SessionCurrency'}))
            currency = currency.split('title="')[1].split('"><div id="SessionCurrency"')[0]
            
            #Parsing Turnover
            turnover = str(soup.find('td', {'class' : 'grid_cell BetSum'}))
            turnover = turnover.split('title="')[1].split('"><div id="BetSum"')[0]
            
            #Parsing Payoff
            payoff = str(soup.find('td', {'class' : 'grid_cell WinSum'}))
            payoff = payoff.split('title="')[1].split('"><div id="WinSum"')[0]
            
            #Calculating Net result + Margin
            net = float(payoff) - float(turnover)
            margin = float(net) / float(turnover) * 100
            
            #Formatting values
            turnover = "{:5,.2f}".format(float(turnover))
            net = "{:5,.2f}".format(net)
            margin = "{:5,.2f}".format(margin) + '%'
            
            #Getting EURO values
            turnover_eur = str(soup.find('td', {'class' : 'grid_cell BetUSD'}))
            turnover_eur = turnover_eur.split('title="')[1].split('"><div id="BetUSD"')[0]
            
            payoff_eur = str(soup.find('td', {'class' : 'grid_cell WinUSD'}))
            payoff_eur = payoff_eur.split('title="')[1].split('"><div id="WinUSD"')[0]
            
            net_eur = float(payoff_eur) - float(turnover_eur)
            
            net_eur = "{:5,.2f}".format(net_eur)
            turnover_eur = "{:5,.2f}".format(float(turnover_eur))

            #Adding to table
            
            cells = table1_inside_cell110.add_row().cells
            cells[0].text = uid
            cells[1].text = screen_name
            cells[2].text = currency + ' ' + turnover
            cells[3].text = currency + ' ' + net
            cells[4].text = margin
            
            if currency != 'EUR':
                cells_eur = table1_inside_cell110.add_row().cells
                cells_eur[2].text = 'EUR ' + turnover_eur
                cells_eur[3].text = 'EUR ' + net_eur

        #Adding Text
        
        cell011.text = operator_id
        
        #1st Paragraph
        
        intro_paragraph = cell110.paragraphs[0]
        
        if report_type_var.get() == 'Players from same casino':
            pass
        else:
            intro = intro.replace('Accounts', 'Account')
            intro = intro.replace('are', 'is')
            intro = intro.replace("players'", "player's")
    
        intro_paragraph.text = intro
        
        #2nd Paragraph
        
        if report_type_var.get() == 'Players from same casino':
            pass
        elif report_type_var.get() == 'Players from different casino':
            explanatory1 = explanatory1.replace(
                'the players below are', 
                'the player below is')
            explanatory1 = explanatory1.replace(
                'accounts\'', 
                'account\'s')        
        elif report_type_var.get() == 'Opposite with himself':
            explanatory1 = explanatory1.replace(
                'the players below are', 
                'the player below is')
            explanatory1 = explanatory1.replace(
                'playing in pairs with each other by ', 
                '')
            explanatory1 = explanatory1.replace(
                'accounts\'', 
                'account\'s')
            explanatory1 = explanatory1.replace(
                'spots', 
                'spots against himself')        
        
        explanatory1_paragraph = cell110.paragraphs[1]
        explanatory1_paragraph.text = explanatory1
        
        #3rd Paragraph
        
        explanatory2_paragraph = cell110.paragraphs[2]
        
        if report_type_var.get() == 'Players from same casino':
            pass
        elif report_type_var.get() == 'Players from different casino':
            explanatory2_roulette = explanatory2_roulette.replace(
                'The players have', 
                'The player with another player from a different casino has')
            explanatory2_baccarat = explanatory2_baccarat.replace(
                'The players have', 
                'The player with another player from a different casino has')
        elif report_type_var.get() == 'Opposite with himself':
            explanatory2_roulette = explanatory2_roulette.replace(
                'The players have', 'The player has')
            explanatory2_baccarat = explanatory2_baccarat.replace(
                'against each other', 'on opposite betting spots on his own')
        
        if game_type_var.get() == 'Roulette':
            explanatory2_paragraph.text = explanatory2_roulette
        else:
            explanatory2_paragraph.text = explanatory2_baccarat
        
        #Conclusion
        
        if report_type_var.get() == 'Players from same casino':
            cell210.text = conclusion
        else:
            conclusion = conclusion.replace('These players', 'This player')
            cell210.text = conclusion
    
    except IndexError:
        generate_text_var.set("WRONG LINK! TRY AGAIN!")
    
    #Save Docx
    
    try:
        document.save(
            'Opposite betting Report ' + operator_id + time.strftime(
                " %d.%m.%y") + '.docx')
    except PermissionError:
        document.save(
            'Opposite betting Report ' + operator_id + time.strftime(
                " %d.%m.%y") + '(1).docx')

#Tkinter Architecture

root = Tk()
root.title('Opposite players Reporting')
root.geometry('750x220')

#Entry 1. Timeframe of analysis

frame1 = Frame(root)
frame1.pack()

dates_label = Label(frame1, 
                    text='Timeframe of analysis (ALL HISTORY by default)')
dates_label.pack(padx=40, side=LEFT)

dates_entry = Entry(frame1)
dates_entry.pack(side=RIGHT)

#Report Type

frame2 = Frame(root)
frame2.pack()

report_type_label = Label(frame2, text='Report Type')
report_type_label.pack(padx=134, pady=7, side=LEFT)

report_type_list = ['Players from same casino', 'Players from different casino', 'Opposite with himself']
report_type_var = StringVar(frame2)
report_type_var.set('Click to choose')

report_type_optionmenu = OptionMenu(frame2, report_type_var, *report_type_list)
report_type_optionmenu.pack(padx=93, side=RIGHT)

#Game Type

frame3 = Frame(root)
frame3.pack()

game_type_label = Label(frame3, text='Game type')
game_type_label.pack(padx=137, side=LEFT)

game_type_list = ['Roulette', 'Baccarat']
game_type_var = StringVar(frame3)
game_type_var.set('Click to choose')

game_type_optionmenu = OptionMenu(frame3, game_type_var, *game_type_list)
game_type_optionmenu.pack(padx=94, side=RIGHT)


#Field to put link from BO with player name and Brand P1

frame4 = Frame(root)
frame4.pack()

link_label_var = StringVar(frame4)
link_label = Label(frame4, 
                   textvariable=link_label_var)
link_label_var.set("BO Link")

link_label.pack(padx=146, pady=7, side=LEFT)

link_entry = Entry(frame4)
link_entry.pack(padx=105, pady=7, side=RIGHT)

#ADD TO LIST

frame_last = Frame(root)
frame_last.pack()

add_text_var = StringVar(frame_last)

add = Button(frame_last, 
                  textvariable=add_text_var)
add_text_var.set("Add Link")

add.pack(fill = X, padx = 130, pady = 10)
add.bind('<Button-1>', list_append)

#Generate

frame_generate = Frame(root)
frame_generate.pack()

generate_text_var = StringVar(frame_generate)

generate = Button(frame_generate, 
                  textvariable=generate_text_var)
generate_text_var.set("Generate")

generate.pack(fill=X, pady=10)
generate.bind('<Button-1>', main_function)

#End

root.mainloop()
